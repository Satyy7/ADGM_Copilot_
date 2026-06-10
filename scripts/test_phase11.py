"""Phase 11 end-to-end test suite.

Run with:
    uv run python scripts/test_phase11.py

Requires:
  - Backend server running on port 8001
  - historical_reviews Qdrant collection seeded (run seed_historical_reviews.py first)
  - A sample DOCX file at scripts/sample_shareholder_resolution.docx (created by Phase 9 tests)

Tests
-----
1. POST /cases/search - basic similarity search
2. POST /cases/search - document-type-specific query
3. POST /cases/search - violation-type query
4. POST /reviews/analyze - verify similar_cases field populated
5. GET  /cases/search - verify 422 on invalid payload
"""

import json
import os
import sys
import urllib.error
import urllib.request

os.environ.setdefault("PYTHONIOENCODING", "utf-8")

BASE = "http://localhost:8001/api/v1"


def post_json(path: str, body: dict, timeout: int = 180) -> dict:
    data = json.dumps(body).encode()
    req = urllib.request.Request(
        f"{BASE}{path}",
        data=data,
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    with urllib.request.urlopen(req, timeout=timeout) as r:
        return json.loads(r.read())


def post_file(path: str, file_path: str, timeout: int = 180) -> dict:
    """Upload a file using multipart/form-data."""
    import email.mime.multipart
    import mimetypes
    import uuid
    boundary = uuid.uuid4().hex
    with open(file_path, "rb") as fh:
        file_bytes = fh.read()
    filename = os.path.basename(file_path)
    mime_type = mimetypes.guess_type(filename)[0] or "application/octet-stream"
    body = (
        f"--{boundary}\r\n"
        f'Content-Disposition: form-data; name="file"; filename="{filename}"\r\n'
        f"Content-Type: {mime_type}\r\n\r\n"
    ).encode() + file_bytes + f"\r\n--{boundary}--\r\n".encode()
    req = urllib.request.Request(
        f"{BASE}{path}",
        data=body,
        headers={"Content-Type": f"multipart/form-data; boundary={boundary}"},
        method="POST",
    )
    with urllib.request.urlopen(req, timeout=timeout) as r:
        return json.loads(r.read())


all_passed = True

# ---------------------------------------------------------------------------
# Test 1: basic similarity search — AoA UBO violations
# ---------------------------------------------------------------------------
print(f"\n{'=' * 68}")
print("TEST 1: POST /cases/search — AoA with UBO violations query")
print(f"{'=' * 68}")
try:
    data = post_json("/cases/search", {
        "query": "articles of association with UBO beneficial ownership violations",
        "top_k": 3,
    })
    cases = data.get("similar_cases", [])
    total = data.get("total_found", 0)
    print(f"Total found: {total}")
    for i, c in enumerate(cases, 1):
        print(f"  [{i}] {c['document_name']} | score={c['compliance_score']} | "
              f"similarity={c['similarity_score']} | violations={c['violation_count']}")
    if total > 0:
        print("PASS")
    else:
        print("FAIL -- no similar cases returned (is the collection seeded?)")
        all_passed = False
except Exception as exc:
    print(f"ERROR: {type(exc).__name__}: {exc}")
    all_passed = False

# ---------------------------------------------------------------------------
# Test 2: employment contract probation query
# ---------------------------------------------------------------------------
print(f"\n{'=' * 68}")
print("TEST 2: POST /cases/search — employment probation violation")
print(f"{'=' * 68}")
try:
    data = post_json("/cases/search", {
        "query": "employment contract probation period exceeds statutory maximum ADGM",
        "top_k": 3,
    })
    cases = data.get("similar_cases", [])
    total = data.get("total_found", 0)
    print(f"Total found: {total}")
    for i, c in enumerate(cases, 1):
        print(f"  [{i}] {c['document_name']} | type={c['document_type']} | "
              f"similarity={c['similarity_score']}")
    # Expect beta_tech_employment_contract to be in top results
    found_employment = any("employment" in c.get("document_type", "") for c in cases)
    if total > 0:
        print("PASS")
        if found_employment:
            print("  Bonus: employment contract case correctly ranked in results")
    else:
        print("FAIL -- no results returned")
        all_passed = False
except Exception as exc:
    print(f"ERROR: {type(exc).__name__}: {exc}")
    all_passed = False

# ---------------------------------------------------------------------------
# Test 3: board resolution governance query
# ---------------------------------------------------------------------------
print(f"\n{'=' * 68}")
print("TEST 3: POST /cases/search — board resolution conflict of interest")
print(f"{'=' * 68}")
try:
    data = post_json("/cases/search", {
        "query": "board resolution conflict of interest director disclosure ADGM governance",
        "top_k": 5,
    })
    cases = data.get("similar_cases", [])
    total = data.get("total_found", 0)
    print(f"Total found: {total}")
    for i, c in enumerate(cases, 1):
        print(f"  [{i}] {c['document_name']} | type={c['document_type']} | "
              f"score={c['compliance_score']} | similarity={c['similarity_score']}")
    if total > 0:
        print("PASS")
    else:
        print("FAIL -- no results returned")
        all_passed = False
except Exception as exc:
    print(f"ERROR: {type(exc).__name__}: {exc}")
    all_passed = False

# ---------------------------------------------------------------------------
# Test 4: compliant document query — expect high-score result
# ---------------------------------------------------------------------------
print(f"\n{'=' * 68}")
print("TEST 4: POST /cases/search — shareholder resolution UBO update")
print(f"{'=' * 68}")
try:
    data = post_json("/cases/search", {
        "query": "shareholder resolution to update UBO register beneficial ownership ADGM",
        "top_k": 3,
    })
    cases = data.get("similar_cases", [])
    total = data.get("total_found", 0)
    print(f"Total found: {total}")
    for i, c in enumerate(cases, 1):
        print(f"  [{i}] {c['document_name']} | score={c['compliance_score']} | "
              f"similarity={c['similarity_score']}")
    if total > 0:
        print("PASS")
    else:
        print("FAIL -- no results returned")
        all_passed = False
except Exception as exc:
    print(f"ERROR: {type(exc).__name__}: {exc}")
    all_passed = False

# ---------------------------------------------------------------------------
# Test 5: invalid payload — expect 422
# ---------------------------------------------------------------------------
print(f"\n{'=' * 68}")
print("TEST 5: POST /cases/search — validation error (query too short)")
print(f"{'=' * 68}")
try:
    post_json("/cases/search", {"query": "hi", "top_k": 3})
    print("FAIL -- expected 422 but got 200")
    all_passed = False
except urllib.error.HTTPError as exc:
    if exc.code == 422:
        print(f"Got HTTP 422 as expected  PASS")
    else:
        print(f"FAIL -- expected 422 but got {exc.code}")
        all_passed = False
except Exception as exc:
    print(f"ERROR: {type(exc).__name__}: {exc}")
    all_passed = False

# ---------------------------------------------------------------------------
# Test 6: POST /reviews/analyze — verify similar_cases in response
#         Uses the same sample DOCX from Phase 9 tests if it exists
# ---------------------------------------------------------------------------
print(f"\n{'=' * 68}")
print("TEST 6: POST /reviews/analyze — similar_cases field in response")
print(f"{'=' * 68}")

SAMPLE_DOC = "scripts/sample_shareholder_resolution.docx"
if not os.path.exists(SAMPLE_DOC):
    print(f"Sample doc not found at {SAMPLE_DOC} -- creating a minimal test DOCX...")
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument()
        doc.add_heading("SHAREHOLDER RESOLUTION", 0)
        doc.add_paragraph(
            "RESOLVED that the company updates its beneficial ownership (UBO) register "
            "in accordance with ADGM Beneficial Ownership Regulations 2018. "
            "The following individuals are confirmed as beneficial owners holding more than 25%: "
            "1. John Smith - 60% ownership. "
            "PASSED by special resolution with 100% shareholder approval."
        )
        doc.save(SAMPLE_DOC)
        print(f"Created minimal test DOCX at {SAMPLE_DOC}")
    except Exception as exc:
        print(f"Could not create test DOCX: {exc} -- skipping test 6")
        SAMPLE_DOC = None

if SAMPLE_DOC and os.path.exists(SAMPLE_DOC):
    try:
        data = post_file("/reviews/analyze", SAMPLE_DOC, timeout=300)
        doc_type = data.get("document_type", "?")
        score    = data.get("compliance_score", "?")
        similar  = data.get("similar_cases", [])
        model    = data.get("model", "?")
        print(f"Document type : {doc_type}")
        print(f"Compliance    : {score}/100")
        print(f"Model         : {model}")
        print(f"Similar cases : {len(similar)}")
        for i, c in enumerate(similar, 1):
            print(f"  [{i}] {c['document_name']} | similarity={c['similarity_score']}")
        if "similar_cases" in data:
            print("PASS -- similar_cases field present in response")
        else:
            print("FAIL -- similar_cases field missing from response")
            all_passed = False
    except urllib.error.HTTPError as exc:
        print(f"HTTP {exc.code}: {exc.read().decode()[:300]}")
        all_passed = False
    except Exception as exc:
        print(f"ERROR: {type(exc).__name__}: {exc}")
        all_passed = False

# ---------------------------------------------------------------------------
# Summary
# ---------------------------------------------------------------------------
print(f"\n{'=' * 68}")
print("ALL TESTS PASSED" if all_passed else "SOME TESTS FAILED")
print(f"{'=' * 68}")
sys.exit(0 if all_passed else 1)
