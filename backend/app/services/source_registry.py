"""Source registry extraction and classification.

`Data Sources.docx` is treated as the seed manifest for official ADGM sources.
This module extracts the table, classifies each URL, and writes a normalized
manifest used by ingestion.
"""

from __future__ import annotations

import hashlib
import json
from datetime import UTC, datetime
from pathlib import Path

from docx import Document

from backend.app.schemas.source import SourceFormat, SourceManifest, SourceRecord


def build_source_id(url: str) -> str:
    """Build a stable source identifier from a URL."""

    digest = hashlib.sha256(url.strip().encode("utf-8")).hexdigest()
    return digest[:16]


def infer_source_format(url: str) -> SourceFormat:
    """Infer source format from the URL path."""

    normalized = url.lower().split("?")[0]
    if ".pdf" in normalized:
        return "pdf"
    if ".docx" in normalized:
        return "docx"
    if normalized.startswith("http"):
        return "html"
    return "unknown"


def infer_collection(category: str, document_type: str, url: str) -> str:
    """Infer the target Qdrant collection for a source row."""

    text = f"{category} {document_type} {url}".lower()
    source_format = infer_source_format(url)
    if "rulebook" in text or "regulatory guidance" in text:
        return "regulations"
    if "checklist" in text:
        return "checklists"
    if source_format == "docx":
        return "templates"
    if "appropriate policy document template" in text:
        return "templates"
    if "policy & guidance" in text or "guidance" in category.lower():
        return "guidance"
    return "guidance"


def extract_source_manifest(docx_path: Path) -> SourceManifest:
    """Extract source rows from the registry DOCX."""

    document = Document(docx_path)
    if not document.tables:
        raise ValueError(f"No source table found in {docx_path}.")

    table = document.tables[0]
    sources: list[SourceRecord] = []
    seen_urls: set[str] = set()

    for row in table.rows[1:]:
        cells = [cell.text.strip() for cell in row.cells]
        if len(cells) < 3:
            continue

        category, document_type, url = cells[:3]
        if not url.startswith("http") or url in seen_urls:
            continue

        seen_urls.add(url)
        sources.append(
            SourceRecord(
                source_id=build_source_id(url),
                category=category,
                document_type=document_type,
                url=url,
                source_format=infer_source_format(url),
                collection=infer_collection(category, document_type, url),
            )
        )

    return SourceManifest(
        generated_at=datetime.now(UTC),
        source_document=str(docx_path),
        sources=sources,
    )


def save_source_manifest(manifest: SourceManifest, output_path: Path) -> None:
    """Write a source manifest as pretty JSON."""

    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(
        manifest.model_dump_json(indent=2),
        encoding="utf-8",
    )


def load_source_manifest(path: Path) -> SourceManifest:
    """Read a source manifest from JSON."""

    return SourceManifest.model_validate(json.loads(path.read_text(encoding="utf-8")))
