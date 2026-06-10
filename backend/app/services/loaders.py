"""Loaders for HTML, PDF, and DOCX official ADGM sources."""

from __future__ import annotations

import hashlib
import html
import re
import urllib.request
from datetime import UTC, datetime
from html.parser import HTMLParser
from pathlib import Path

import fitz
from docx import Document

from backend.app.schemas.source import (
    NormalizedDocument,
    NormalizedSection,
    SourceRecord,
)


class HTMLSectionParser(HTMLParser):
    """Extract readable heading-aware text sections from an HTML document."""

    ignored_tags = {"script", "style", "noscript", "svg", "nav", "footer"}
    heading_tags = {"h1", "h2", "h3", "h4"}
    block_tags = {
        "p",
        "li",
        "td",
        "th",
        "blockquote",
        "article",
        "section",
        "div",
    }

    def __init__(self) -> None:
        """Initialize parser state."""

        super().__init__()
        self._ignore_depth = 0
        self._current_tag: str | None = None
        self._buffer: list[str] = []
        self._heading_path: list[str] = []
        self.sections: list[NormalizedSection] = []
        self.links: list[str] = []
        self.title: str | None = None

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        """Track relevant tags and links."""

        if tag in self.ignored_tags:
            self._ignore_depth += 1
            return
        if self._ignore_depth:
            return
        if tag == "a":
            href = dict(attrs).get("href")
            if href:
                self.links.append(href)
        if tag in self.heading_tags or tag in self.block_tags or tag == "title":
            self.flush()
            self._current_tag = tag

    def handle_endtag(self, tag: str) -> None:
        """Flush text when a relevant block closes."""

        if tag in self.ignored_tags and self._ignore_depth:
            self._ignore_depth -= 1
            return
        if self._ignore_depth:
            return
        if tag == self._current_tag:
            self.flush()
            self._current_tag = None

    def handle_data(self, data: str) -> None:
        """Collect readable text."""

        if self._ignore_depth:
            return
        cleaned = normalize_whitespace(html.unescape(data))
        if cleaned:
            self._buffer.append(cleaned)

    def flush(self) -> None:
        """Convert buffered text into a heading or section."""

        text = normalize_whitespace(" ".join(self._buffer))
        self._buffer = []
        if not text:
            return

        tag = self._current_tag
        if tag == "title":
            self.title = text
            return
        if tag in self.heading_tags:
            level = int(tag[1])
            self._heading_path = self._heading_path[: level - 1]
            self._heading_path.append(text)
            return

        self.sections.append(
            NormalizedSection(
                heading=self._heading_path[-1] if self._heading_path else None,
                section_path=list(self._heading_path),
                text=text,
                links=list(self.links),
            )
        )


def normalize_whitespace(value: str) -> str:
    """Collapse repeated whitespace while preserving readable text."""

    return re.sub(r"\s+", " ", value).strip()


def file_sha256(path: Path) -> str:
    """Return the SHA-256 hash for a file."""

    digest = hashlib.sha256()
    with path.open("rb") as file:
        for block in iter(lambda: file.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def raw_file_extension(source: SourceRecord) -> str:
    """Return the expected raw file extension for a source."""

    if source.source_format == "pdf":
        return ".pdf"
    if source.source_format == "docx":
        return ".docx"
    return ".html"


def download_source(source: SourceRecord, raw_dir: Path, timeout_seconds: int = 30) -> Path:
    """Download one source into the raw data directory."""

    target_dir = raw_dir / source.collection
    target_dir.mkdir(parents=True, exist_ok=True)
    target_path = target_dir / f"{source.source_id}{raw_file_extension(source)}"
    request = urllib.request.Request(
        str(source.url),
        headers={
            "User-Agent": (
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 (KHTML, like Gecko) "
                "Chrome/125.0 Safari/537.36 ADGM-Compliance-Copilot/0.1"
            ),
            "Accept": (
                "text/html,application/xhtml+xml,application/xml;q=0.9,"
                "application/pdf,application/vnd.openxmlformats-"
                "officedocument.wordprocessingml.document,*/*;q=0.8"
            ),
            "Accept-Language": "en-US,en;q=0.9",
        },
    )
    with urllib.request.urlopen(request, timeout=timeout_seconds) as response:
        target_path.write_bytes(response.read())
    return target_path


def load_html(source: SourceRecord, path: Path) -> NormalizedDocument:
    """Load and normalize an HTML source."""

    parser = HTMLSectionParser()
    raw_html = path.read_text(encoding="utf-8", errors="ignore")
    parser.feed(extract_relevant_html(raw_html))
    parser.flush()
    title = parser.title or source.document_type
    sections = merge_short_sections(filter_boilerplate_sections(parser.sections))
    return NormalizedDocument(
        source=source,
        title=title,
        downloaded_path=str(path),
        file_hash=file_sha256(path),
        extracted_at=datetime.now(UTC),
        sections=sections,
    )


def extract_relevant_html(raw_html: str) -> str:
    """Focus extraction on page-specific content blocks when possible."""

    content_match = re.search(
        r"<content[^>]*>(.*?)</content>",
        raw_html,
        flags=re.IGNORECASE | re.DOTALL,
    )
    if content_match:
        return content_match.group(1)

    main_match = re.search(
        r"<main[^>]*>(.*?)</main>",
        raw_html,
        flags=re.IGNORECASE | re.DOTALL,
    )
    if main_match:
        return main_match.group(1)

    adgm_body_blocks = re.findall(
        r"<adgm-body[^>]*>(.*?)</adgm-body>",
        raw_html,
        flags=re.IGNORECASE | re.DOTALL,
    )
    if adgm_body_blocks:
        title_match = re.search(
            r"<h1[^>]*>(.*?)</h1>",
            raw_html,
            flags=re.IGNORECASE | re.DOTALL,
        )
        title_html = title_match.group(0) if title_match else ""
        return title_html + "\n".join(adgm_body_blocks)

    return raw_html


def filter_boilerplate_sections(
    sections: list[NormalizedSection],
) -> list[NormalizedSection]:
    """Remove common navigation and site chrome sections."""

    return [section for section in sections if not is_boilerplate_text(section.text)]


def is_boilerplate_text(text: str) -> bool:
    """Return True for common ADGM site navigation/marketing chrome."""

    normalized = text.lower()
    boilerplate_phrases = (
        "skip to main content",
        "overview jurisdiction authorities initiatives",
        "a platform that offers limitless opportunities",
        "read our latest news articles",
        "follow us",
        "terms and conditions",
        "privacy site map",
        "all rights reserved",
        "about adgm legal framework public registers faqs lifestyle",
        "adgm academy dispute resolution careers",
    )
    if any(phrase in normalized for phrase in boilerplate_phrases):
        return True
    if len(text) < 25 and normalized in {"print", "search", "custom print"}:
        return True
    return False


def load_pdf(source: SourceRecord, path: Path) -> NormalizedDocument:
    """Load and normalize a PDF source using PyMuPDF."""

    sections: list[NormalizedSection] = []
    with fitz.open(path) as pdf:
        title = pdf.metadata.get("title") or source.document_type
        for page_index, page in enumerate(pdf, start=1):
            text = normalize_whitespace(page.get_text("text"))
            if text:
                sections.append(
                    NormalizedSection(
                        heading=f"Page {page_index}",
                        section_path=[source.document_type, f"Page {page_index}"],
                        text=text,
                        page_number=page_index,
                    )
                )
    return NormalizedDocument(
        source=source,
        title=title,
        downloaded_path=str(path),
        file_hash=file_sha256(path),
        extracted_at=datetime.now(UTC),
        sections=sections,
    )


def load_docx(source: SourceRecord, path: Path) -> NormalizedDocument:
    """Load and normalize a DOCX source."""

    document = Document(path)
    sections: list[NormalizedSection] = []
    heading_path: list[str] = []
    paragraph_buffer: list[str] = []

    def flush_paragraphs() -> None:
        text = normalize_whitespace(" ".join(paragraph_buffer))
        paragraph_buffer.clear()
        if text:
            sections.append(
                NormalizedSection(
                    heading=heading_path[-1] if heading_path else None,
                    section_path=list(heading_path),
                    text=text,
                )
            )

    for paragraph in document.paragraphs:
        text = normalize_whitespace(paragraph.text)
        if not text:
            continue
        style_name = (paragraph.style.name or "").lower()
        if "heading" in style_name or text.isupper():
            flush_paragraphs()
            heading_path = [text]
        else:
            paragraph_buffer.append(text)
    flush_paragraphs()

    for table_index, table in enumerate(document.tables, start=1):
        rows = [
            [normalize_whitespace(cell.text) for cell in row.cells]
            for row in table.rows
        ]
        table_text = " ".join(" | ".join(row) for row in rows)
        if table_text.strip():
            sections.append(
                NormalizedSection(
                    heading=f"Table {table_index}",
                    section_path=[source.document_type, f"Table {table_index}"],
                    text=normalize_whitespace(table_text),
                    tables=[rows],
                )
            )

    return NormalizedDocument(
        source=source,
        title=source.document_type,
        downloaded_path=str(path),
        file_hash=file_sha256(path),
        extracted_at=datetime.now(UTC),
        sections=merge_short_sections(sections),
    )


def load_source_file(source: SourceRecord, path: Path) -> NormalizedDocument:
    """Dispatch a raw source file to the correct loader."""

    if source.source_format == "pdf":
        return load_pdf(source, path)
    if source.source_format == "docx":
        return load_docx(source, path)
    if source.source_format == "html":
        return load_html(source, path)
    raise ValueError(f"Unsupported source format: {source.source_format}")


def merge_short_sections(
    sections: list[NormalizedSection],
    minimum_chars: int = 220,
) -> list[NormalizedSection]:
    """Merge tiny adjacent sections with the same heading path."""

    merged: list[NormalizedSection] = []
    for section in sections:
        if (
            merged
            and len(merged[-1].text) < minimum_chars
            and merged[-1].section_path == section.section_path
            and merged[-1].page_number == section.page_number
        ):
            merged[-1].text = normalize_whitespace(f"{merged[-1].text} {section.text}")
            merged[-1].links.extend(section.links)
        else:
            merged.append(section)
    return merged
